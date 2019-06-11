from docx import Document
from docx.shared import Inches, Pt, RGBColor, Length
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.dml.color import ColorFormat
from docx.text.run import Font, Run
from docx.enum.text import WD_ALIGN_PARAGRAPH #does not work for ryans pycharm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #this one works for ryans pycharm
from docx.oxml.shared import OxmlElement, qn
from copy import deepcopy
import datetime
import textwrap

#pip3 install python-docx
#pip3 install DateTime

# This class takes a template file and then fills in the correct data. It then will
# get emailed for approval.
# template file must be in the same directory or its location specified.
# To instance form gen one should provide the correct OC insturctor name, OC department,
# military_course, and the oc_course

#creating a FileGen object will fill out the top of the form.

#----- About: Nested_Tables_Form_Template_Version_4_11_2019-----
#If things dont line up you have to make sure that the margin for the cell where the mlt outcomes go is lined up 
#so that the courser is just showing. If it still doesnt line up then you will have to mess with it. But I am hoping
#that the formating follows the form template.

class FileGen:

        # initializing variables to be used throught the generator
    def __init__(self, oc_course, jst_course, reviewer):
        # test.docx should be the name of the template file.
        self.doc = Document('Nested_Tables_Form_Template_Version_4_11_2019')# Have to use this form template.
        self.total_tables = self.doc.tables  # locating document tables
        # locating correct table for course comparason
        self.info_table = self.total_tables[0]
        self.date_cell = self.info_table.cell(0, 0)
        self.military_course_cell = self.info_table.cell(0, 1)
        # data for filling in top of doc.
        self.instructor_name_cell = self.info_table.cell(1, 0)
        self.oc_course_cell = self.info_table.cell(1, 1)
        self.instr_name = reviewer.name
        self.dep_name = reviewer.department
        self.m_course = jst_course.number
        self.oc_course = oc_course.number
        self.oc_description = oc_course.description
        self.mc_description = jst_course.description
        self.oc_course_name = oc_course.name
        self.mc_course_name = jst_course.name
        # bellow marks out the columns for both military and olivet course
        # outcomes.
        self.oc_column = 0
        self.total_columns_added = 1 # this should actually be total rows added. will change eventualy.
        self.total_oc_course = 0
        self.total_columns = 6
        # used for sugestive guesses 
        self.no_match = 33
        self.moderate_match = 67
        self.strong_match = 100
        self.line_spacing = 12
        self.remove_last_row = False
        
        self.width_info = []
        self.table_iterator = 0
        self.ran_find_split_and_copy = False
        self.nested_mc_tables = []
        self.mc_column = 0
        self.compare_tables, self.first_comp_row = self.Find_Compare_Tables()
        self.comp_table = self.compare_tables[0]
        self.nested_mc_tables.append(self.Create_Nested_Table(self.comp_table))
        self.Update_Compare_Info()
        self.mlt_table = self.nested_mc_tables[self.table_iterator]
        self.mlt_row = 0
        self.mlt_column = 0
        self.mlt_rows = self.mlt_table.rows
        self.Check_Table_Style()
        self.Fill_Course_Info()

    def Update_Compare_Info(self, compare_table=None):
        self.compare_tables, self.first_comp_row = self.Find_Compare_Tables() # finding the first compare row and compare table, results could difer per form.
        self.comp_table = self.compare_tables[self.table_iterator]
        self.mlt_table = self.nested_mc_tables[self.table_iterator]
        self.comp_rows = self.comp_table.rows
        self.comp_last_row = self.comp_rows[-1]
        self.mlt_rows = self.mlt_table.rows
        self.comp_copy_row = self.comp_rows[1]
        self.mc_outcome_cell = self.comp_table.cell(1,1)
        self.comp_row = self.first_comp_row
        self.tbl = self.comp_table._tbl
        self.mc_row = len(self.comp_rows)-1
        self.mc_column = 1
        self.oc_row = 3
        
    def Find_Compare_Tables(self):
        tables = []
        f_comp_row = 0
        iterator = 0
        for table in self.total_tables:
            iterator = 1
            for row in table.rows:
                for cell in row.cells:
                    if cell.paragraphs[0].text == "Olivet College\nCourse Learning Outcome":
                        tables.append(table)
                        f_comp_row = iterator
                iterator += 1
        return tables, f_comp_row

    # called during initualization fills in the top of the table.
    def Fill_Course_Info(self):
        now = datetime.datetime.now()
        self.date_cell.text = "Date of Initiation:\n" + \
            str(now.month) + "-" + str(now.day) + "-" + str(now.year)
        self.military_course_cell.text = "JST or AU course for which Credit/Equivalency is sought:\n" + self.m_course + " " + self.mc_course_name + "\n" + self.mc_description
        self.instructor_name_cell.text = "Evaluator Name:\n" + \
            self.instr_name + "\nDepartment:\n" + self.dep_name
        self.oc_course_cell.text = "Olivet College course being considered for possible equivalency:\n" + self.oc_course + " " + self.oc_course_name + "\n" + self.oc_description


    def GetLastRow(self, rows):
        last = rows[-1]
        return last

    def Check_Table_Style(self):
        if self.comp_table.style != 'Table Grid':
            self.comp_table.style = 'Table Grid'

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def copy_table_after(self, ctable, atable):
        ctbl, atbl = ctable._tbl, atable._tbl
        new_tbl = deepcopy(ctbl)
        atbl.addnext(new_tbl)

    def move_para_after(self, paragraph, table):
        p = paragraph._p
        table.addnext(p)

    def update_tables(self): # used to update the total_tables property
        self.total_tables = self.doc.tables

    # tables should be total number of oc_outcomes.
    def Find_Split_And_Copy(self, tables): # designed to create as many tables as there are oc outcomes.
        self.ran_find_split_and_copy = True
        iterator = tables-1
        for x in range(0,iterator):# goes through and makes the enough tables to fit each oc_outcome.
            new_tbl = self.copy_table_after(self.comp_table, self.comp_table)
            iterator -= 1
        self.update_tables()
        self.Update_Compare_Info()
        self.Create_Nested_Tables()

    # checks to see if tables have been generated.
    def Check_Find_Split_And_Copy(self):
        if self.ran_find_split_and_copy == False:
            raise HaveToCreateTables("Have to call Find_Split_And_Copy method first") # Failing to create tables will throw error

    def Remove_Border_Last_Row(self, border=0):
        new_tbl = self.mlt_table._tbl
        rows = new_tbl.getchildren()
        last_row = self.GetLastRow(rows)
        self.Remove_Border(last_row, border=border, outside=True)

    def Olivet_Course_Outcomes(self, c_outcome):
        self.total_oc_course +=1
        compare_table = self.compare_tables[self.table_iterator]
        self.compare_row = compare_table.rows[1]
        self.oc_outcome_cell = self.compare_row.cells[0]
        para = self.oc_outcome_cell.paragraphs[0]
        para.add_run(c_outcome)

    def JST_Outcomes(self, create_row, jst_outcome, percent, new_row=False):
        self.compare_row = self.mlt_rows[self.mlt_row]
        self.mc_outcome_cell = self.compare_row.cells[self.mlt_column]
        para = self.mc_outcome_cell.paragraphs[0]
        para.add_run(jst_outcome)
        para.paragraph_format.line_spacing = None
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if create_row == True:
            if new_row==False:
                self.Add_Row_At(self.total_columns_added)
            else:
                self.Add_Row_At(self.total_columns_added, border=0)
            self.mlt_row += 1
        if create_row == False and len(self.compare_tables) == 1:
            self.remove_last_row = True
            self.Add_Row_At(self.total_columns_added, border=0)
            self.mlt_row += 1
        self.Add_Checkbox(jst_outcome, percent)

    # a percentage determined by nltk to highlight a sugested box to check.
    def Sugested_Check(self, percent, font, row, para, run):
        if(percent >= 1 and percent <= self.no_match and row == 1):
            #self.__Highlight(font)
            self.Check_Mark(run)
        if( percent > self.no_match and percent <= self.moderate_match and row == 2):
            #self.__Highlight(font)
            self.Check_Mark(run)
        if(percent > self.moderate_match and percent <= self.strong_match and row == 3):
            #self.__Highlight(font)
            self.Check_Mark(run)
        para.paragraph_format.line_spacing = Pt(self.line_spacing)

    def Add_Checkbox(self, jst_outcome, percent):
            column_check_add = self.mlt_row

            for i in range(self.mlt_column+1, len(self.mlt_table.columns)):
                cell_check_add = self.compare_row.cells[i]
                para = cell_check_add.paragraphs[0]
                run = para.add_run("\u2610")
                font = run.font
                self.Sugested_Check(percent, font, i, para, run)
                para.paragraph_format.line_spacing = Pt(self.line_spacing)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #used to generate more rows for the use of comparason. first adds row to the bottom of the table
    #and then moves it to the correct location.
    def Add_Row_At(self, location, border=0):
        new_row = self.mlt_table.add_row()
        tr = new_row._tr
        self.Remove_Border_Last_Row(border=border)# the border is removed based on where it will be incerted into the table.
        if location != "end":# this gives the method the ability to add rows to the end of the table.
            self.tbl.insert(5 + location, tr)#6 meens insert it at the 4th row. 7-5 e.t.c.
            self.total_columns_added += 1

    # this is called after find_split_and_copy() creats the correct tables
    # !!make sure you are useing the correct format when using this method!!
    def Like_Outcome_Tables(self, c_outcome, jst_outcome):
        self.Check_Find_Split_And_Copy()# Checking to see if the tables where generated.
        self.Update_Compare_Info()
        #self.Update_Compare_Info(self.compare_tables[self.table_iterator], self.nested_mc_tables[self.table_iterator])# used to move on to the next table after one is completed.
        self.total_columns_added = "end" # each comparison now has its own table so we down have to incert rows at a point.
        self.Like_Outcomes(c_outcome, jst_outcome)
        self.table_iterator += 1
        self.comp_row = 1

    def Like_Outcomes(self, c_outcome, jst_outcome):           
        self.Olivet_Course_Outcomes(c_outcome)
        iterator = 1
        for outcome in jst_outcome:
            if iterator < len(jst_outcome):
                if len(jst_outcome) == iterator+1:
                    self.JST_Outcomes(True, outcome[0], outcome[1], new_row=True)
                else:
                    self.JST_Outcomes(True, outcome[0], outcome[1])
            else:
                self.JST_Outcomes(False, outcome[0], outcome[1], new_row=True)
            iterator +=1
        self.mlt_row = 0

    def Remove_Border(self, row, border=0, outside=False):
        for cell in row:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')
        
            
        
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            if border == 1:
                tcBorders.append(top)
            if border == 2:
                tcBorders.append(bottom)
            if border == 0:
                tcBorders.append(top)
                tcBorders.append(bottom)
            if border == 3:
                tcBorders.append(right)
                tcBorders.append(left)
                tcBorders.append(top)
                tcBorders.append(bottom)
            tcPr.append(tcBorders)
        index = 0
        if outside == True:
            for cell in row:
                tcPr = cell.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'nil')

                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'nil')
                if index == 0:
                    tcBorders.append(left)
                if row.index(cell) == row.index(row[-1]):
                    tcBorders.append(right)
                tcPr.append(tcBorders)
                index += 1

    def get_table_info(self, table):
        row = table.rows
        cell_widths = []
        for cell in range(1, 6):
            cell_widths.append(row[0].cells[cell].width)
        return cell_widths

    def Create_Nested_Tables(self):
        index = 0
        for table in self.compare_tables:
            if index > 0:
                new_table = self.Create_Nested_Table(table)
                self.nested_mc_tables.append(new_table)
                
            index += 1
    
    def Remove_Row(self, row):
        tr = row._tr
        self.tbl.remove(tr)

    def Remove_Empty_Row(self):
        empty_row = self.comp_rows[-4]
        self.Remove_Row(empty_row)

    # used for highlighting runs.
    def Highlight(self, font):
        font.highlight_color = WD_COLOR_INDEX.GRAY_25

    # adds a checkmark for sugestive checking.
    def Check_Mark(self, run):
        run.clear()
        font = run.font
        font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        font.color.rgb = RGBColor(211,211,211)# makes the checkmards light gray
        run.add_text("\u2713")# this is the unicode for checkmark symbol.

    def Create_Nested_Table(self, table):
        table = table
        cell = table.cell(1,1)
        cell._element.clear_content()
        new_table = cell.add_table(1,5)
        new_tbl = new_table._tbl
        self.Remove_Border(new_tbl.getchildren()[2], border=0, outside=True)
        new_table.style = 'Table Grid'
        new_row = new_table.rows[0]

        self.width_info = self.get_table_info(table)
        self.fix_cell_alignment(new_table, self.width_info)
        new_table.autofit = False
        return new_table

    def fix_cell_alignment(self, table, width_info):
        for row in range(0,len(table.rows)):
            for column in range(0,len(table.columns)):
                edit_cell = table.cell(row, column)
                edit_cell.width = width_info[column]

    # will be used if we decide to implement emailing the form.
    def Email_Doc(self):
        pass

    def Save_Doc(self, doc_name='Test-Saved.docx'):
        if self.remove_last_row == True:
            self.Remove_Empty_Row()
        for table in self.nested_mc_tables:
            self.fix_cell_alignment(table, self.width_info)
        self.doc.save(doc_name)
        
# this class is for throwing a error if the like_outcome_table() method is used before the tables are generated using find_split_and_copy()
class HaveToCreateTables(Exception):
    

    def __init__(self, message):
        super().__init__(message)
        self.message = message
        
                

                


                
        