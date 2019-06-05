from docx import Document
from docx.shared import Inches, Pt, RGBColor, Length
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.dml.color import ColorFormat
from docx.text.run import Font, Run
from docx.enum.text import WD_ALIGN_PARAGRAPH #does not work for ryans pycharm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #this one works for ryans pycharm
from docx.oxml.shared import OxmlElement, qn
import datetime
import textwrap

#pip3 install python-docx
#pip3 install DateTime

# This class takes a template file and then fills in the correct data. It then will
# get emailed for approval.
# template file must be in the same directory or its location specified.
# To instance form gen one should provide the correct OC insturctor name, OC department,
# military_course, and the oc_course

#creating a FileGen object will fill out the top of the form.

class FileGen:

        # initializing variables to be used throught the generator
    def __init__(self, instructor_name, department, military_course, oc_course, oc_course_name, mc_course_name, oc_description, mc_description):
        # test.docx should be the name of the template file.
        self.doc = Document('Form_Template_Version_5_13_2019.docx')# Form_Template_Version_5_13_2019.docx # Form_Template_Version_4_11_2019.docx
        self.total_tables = self.doc.tables  # locating document tables
        # locating correct table for course comparason
        self.comp_table = self.total_tables[0]
        self.comp_rows = self.comp_table.rows
        self.comp_copy_row = self.comp_rows[3]
        self.mc_outcome_cell = self.comp_table.cell(3,1)
        #
        self.date_cell = self.comp_table.cell(0, 0)
        self.military_course_cell = self.comp_table.cell(0, 1)
        # data for filling in top of doc.
        self.instructor_name_cell = self.comp_table.cell(1, 0)
        self.oc_course_cell = self.comp_table.cell(1, 1)
        self.instr_name = instructor_name
        self.dep_name = department
        self.m_course = military_course
        self.oc_course = oc_course
        self.oc_description = oc_description
        self.mc_description = mc_description
        self.oc_course_name = oc_course_name
        self.mc_course_name = mc_course_name
        # bellow marks out the columns for both military and olivet course
        # outcomes.
        self.comp_row = 3
        self.mc_row = 3
        self.mc_column = 1
        self.oc_row = 3
        self.oc_column = 0
        self.total_columns_added = 1
        self.total_oc_course = 0
        self.total_columns = 6
        # used for sugestive guesses 
        self.no_match = 33
        self.moderate_match = 67
        self.strong_match = 100
        self.line_spacing = 12
        self.tbl = self.comp_table._tbl

        self.__Remove_Border(self.tbl.getchildren()[5])
        self.__Check_Table_Style()
        self.__Fill_Course_Info()

    # if the table style is different, added rows might not have borders or different borders.
    def __Check_Table_Style(self):
        if self.comp_table.style != 'Table Grid':
            self.comp_table.style = 'Table Grid'

    #used to generate more rows for the use of comparason. first adds row to the bottom of the table
    #and then moves it to the correct location.
    def __Add_Row(self, location, border=0):
        new_row = self.comp_table.add_row()
        tr = new_row._tr
        self.__Remove_Border_Last_Row(border=border)# the border is removed based on where it will be incerted into the table.
        self.tbl.insert(5 + location, tr)#6 meens insert it at the 4th row. 7-5 e.t.c.
        self.total_columns_added += 1
        
    # used in __Remove_Border_Last_Row() to find the last row that is created in __Add_Row()
    def __GetLastRow(self, rows):
        last = rows[-1]
        return last

    # used to remove the borders of cells, either top or bottom.
    def __Remove_Border(self, row, border=0):
        for cell in row:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')
        
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')
        
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            if border == 1:
                tcBorders.append(top)
            if border == 2:
                tcBorders.append(bottom)
            if border == 0:
                tcBorders.append(top)
                tcBorders.append(bottom)
            tcPr.append(tcBorders)

    # the border is removed before it is incerted into to correct spot in the table.
    def __Remove_Border_Last_Row(self, border=0):
        self.compare_row = self.comp_rows[self.comp_row]
        rows = self.tbl.getchildren()
        last_row = self.__GetLastRow(rows)
        self.__Remove_Border(last_row, border=border)

    # used for deleting rows.
    def __Remove_Row(self, row):
        tr = row._tr
        self.tbl.remove(tr)

    # called before saving the program in self.Save_Doc, this function removes the last un-needed row.
    def __Remove_Empty_Row(self):
        empty_row = self.comp_rows[-4]
        self.__Remove_Row(empty_row)

               
    # will add checkboxes the the correct columns
    def __Add_Checkbox(self, jst_outcome, percent):
        column_check_add = self.comp_row

        for i in range(2, len(self.comp_table.columns)):
            cell_check_add = self.compare_row.cells[i]
            para = cell_check_add.paragraphs[0]
            run = para.add_run("\u2610")
            font = run.font
            self.__Sugested_Check(percent, font, i, para, run)
            para.paragraph_format.line_spacing = Pt(self.line_spacing)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #if WD_ALIGN_PARAGRAPH doesnt work for you switch to WD_PARAGRAPH_ALIGNMENT

    # a percentage determined by nltk to highlight a sugested box to check.
    def __Sugested_Check(self, percent, font, row, para, run):
        if(percent >= 1 and percent <= self.no_match and row == 2):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        if( percent > self.no_match and percent <= self.moderate_match and row == 3):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        if(percent > self.moderate_match and percent <= self.strong_match and row == 4):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        para.paragraph_format.line_spacing = Pt(self.line_spacing)

    # adds a checkmark for sugestive checking.
    def __Check_Mark(self, run):
        run.clear()
        font = run.font
        font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        font.color.rgb = RGBColor(211,211,211)# makes the checkmards light gray
        run.add_text("\u2713")# this is the unicode for checkmark symbol.

    # used for highlighting runs.
    def __Highlight(self, font):
        font.highlight_color = WD_COLOR_INDEX.GRAY_25

    # called during initualization fills in the top of the table.
    def __Fill_Course_Info(self):
        now = datetime.datetime.now()
        self.date_cell.text = "Date of Initiation:\n" + \
            str(now.month) + "-" + str(now.day) + "-" + str(now.year)
        self.military_course_cell.text = "JST or AU course for which Credit/Equivalency is sought:\n" + self.m_course + " " + self.mc_course_name + "\n" + self.mc_description
        self.instructor_name_cell.text = "Evaluator Name:\n" + \
            self.instr_name + "\nDepartment:\n" + self.dep_name
        self.oc_course_cell.text = "Olivet College course being considered for possible equivalency:\n" + self.oc_course + " " + self.oc_course_name + "\n" + self.oc_description

    # used for entering individual Olivet course outcomes
    def Olivet_Course_Outcomes(self, c_outcome):
        self.total_oc_course +=1
        self.compare_row = self.comp_rows[self.comp_row]
        self.oc_outcome_cell = self.compare_row.cells[0]
        para = self.oc_outcome_cell.paragraphs[0]
        para.add_run(c_outcome)

    # used for entering individual Military course outcomes.
    # used when the user wants to move to the next cell.
    def JST_Outcomes(self, jst_outcome, percent, new_row=False):
        self.compare_row = self.comp_rows[self.comp_row]
        self.mc_outcome_cell = self.compare_row.cells[1]
        para = self.mc_outcome_cell.paragraphs[0]
        para.add_run(jst_outcome)
        para.paragraph_format.line_spacing = None
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        if new_row==False:
            self.__Add_Row(self.total_columns_added)
        else:
            self.__Add_Row(self.total_columns_added, border=2)
        self.comp_row += 1

        self.__Add_Checkbox(jst_outcome, percent)

    # adds both the Olivet course outcomes with there coresponding military course outcomes.
    # c_outcomes would be just a string for the Olivet college outcome and then jst_outcomes would
    # be the coresponding dictionary that holds percentages.
    def Like_Outcomes(self, c_outcome, jst_outcome):
        self.Olivet_Course_Outcomes(c_outcome)
        iterator = 0
        for outcome in jst_outcome:
            if iterator == len(jst_outcome) - 1:
                self.JST_Outcomes(outcome[0], outcome[1], new_row=True)
            else:
                self.JST_Outcomes(outcome[0], outcome[1])
            iterator += 1


    # will be used if we decide to implement emailing the form.
    def Email_Doc(self):
        pass

    # used to save the document. Must call this to save document.
    def Save_Doc(self, doc_name='Test-Saved.docx'):
        self.__Remove_Empty_Row()# removes the emtpy row before it is saved to the file.
        self.doc.save(doc_name)
