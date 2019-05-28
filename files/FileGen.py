from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.dml.color import ColorFormat
from docx.text.run import Font, Run
from docx.enum.text import WD_ALIGN_PARAGRAPH #does not work for ryans pycharm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #this one works for ryans pycharm
import datetime

#pip3 install python-docx
#pip3 install DateTime

# This class takes a template file and then fills in the correct data. It then will
# get emailed for approval.
# template file must be in the same directory or its location specified.
# To instance form gen one should provide the correct OC insturctor name, OC department,
# military_course, and the oc_course

#creating a FileGen object will fill out the top of the 
class FileGen:

        # initializing variables to be used throught the generator
    def __init__(self, instructor_name, department, military_course, oc_course):
        # test.docx should be the name of the template file.
        self.doc = Document('Test.docx')
        self.total_tables = self.doc.tables  # locating document tables
        # locating correct table for course comparason
        self.comp_table = self.total_tables[0]
        self.comp_rows = self.comp_table.rows
        self.comp_copy_row = self.comp_rows[3]
        #
        self.date_cell = self.comp_table.cell(0, 0)
        self.military_course_cell = self.comp_table.cell(0, 1)
        # data for filling in top of doc.
        self.instructor_name_cell = self.comp_table.cell(1, 0)
        self.oc_course_cell = self.comp_table.cell(1, 1)
        self.instr_name = instructor_name
        self.dep_name = department
        self.m_course = military_course
        self.oc_course = oc_course
        # bellow marks out the columns for both military and olivet course
        # outcomes.
        self.mc_row = 3
        self.mc_column = 1
        self.oc_row = 3
        self.oc_column = 0
        self.total_oc_course = 0
        self.total_columns = 6
        # used for sugestive guesses 
        self.no_match = 33
        self.moderate_match = 67
        self.strong_match = 100
        
        self.line_spacing = 12
        self.__Fill_Course_Info()


    #used to generate more rows at a given point.
    def __Add_Row(self):
        print("creating new row")
        new_row = self.comp_table.add_row()
        tbl = self.comp_table._tbl
        tr = new_row._tr
        print(self.total_oc_course)
        tbl.insert(5 + self.total_oc_course, tr)
               
    # will add checkboxes the the correct columns
    def __Add_Checkbox(self, jst_outcome, percent):
        column_check_add = self.mc_column + 1

        for i in range(column_check_add, len(self.comp_table.columns)):
            cell_check_add = self.comp_table.cell(self.mc_row, i)
            para = cell_check_add.paragraphs[0]
            run = para.add_run("\u2610")
            font = run.font
            self.__Sugested_Check(percent, font, i, para, run)
            if len(jst_outcome) > 50:
                para.add_run("\n\n\n")
            else:
                para.add_run("\n\n")
            para.paragraph_format.line_spacing = Pt(self.line_spacing)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #if WD_ALIGN_PARAGRAPH doesnt work for you switch to WD_PARAGRAPH_ALIGNMENT


    # a percentage determined by nltk to highlight a sugested box to check.
    def __Sugested_Check(self, percent, font, row, para, run):
        if(percent >= 1 and percent <= self.no_match and row == 2):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        if( percent > self.no_match and percent <= self.moderate_match and row == 3):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        if(percent > self.moderate_match and percent <= self.strong_match and row == 4):
            #self.__Highlight(font)
            self.__Check_Mark(run)
        para.paragraph_format.line_spacing = None

    # adds a checkmark for sugestive checking.
    def __Check_Mark(self, run):
        run.clear()
        font = run.font
        print("font is: ", font)
        font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        font.color.rgb = RGBColor(211,211,211)# makes the checkmards light gray
        run.add_text("\u2713")#this is the unicode for checkmark symbol.

    # used for highlighting runs.
    def __Highlight(self, font):
        font.highlight_color = WD_COLOR_INDEX.GRAY_25

    # called during initualization fills in the top of the table.
    def __Fill_Course_Info(self):
        now = datetime.datetime.now()
        self.date_cell.text = "Date of Initiation:\n" + \
            str(now.month) + "-" + str(now.day) + "-" + str(now.year)
        self.military_course_cell.text = "JST or AU course for which Credit/Equivalency is sought:\n" + self.m_course
        self.instructor_name_cell.text = "Evaluator Name:\n" + \
            self.instr_name + "\nDepartment:\n" + self.dep_name
        self.oc_course_cell.text = "Olivet College course being considered for possible equivalency:\n" + self.oc_course

    # used for entering individual Olivet course outcomes
    def Olivet_Course_Outcomes(self, c_outcome):
        if self.total_oc_course >= 1:
            self.__Add_Row()
        self.total_oc_course +=1
        self.oc_outcome_cell = self.comp_table.cell(
            self.oc_row, self.oc_column)
        para = self.oc_outcome_cell.paragraphs[0]
        para.add_run(c_outcome)
        self.oc_row += 1

    # used for entering individual Military course outcomes
    # used when the user wants to move to the next cell.
    def JST_Outcomes(self, jst_outcome, percent, new_cell=False):
        if new_cell == True:
            self.mc_row += 1
        self.mc_outcome_cell = self.comp_table.cell(
            self.mc_row, self.mc_column)
        para = self.mc_outcome_cell.paragraphs[0]
        para.paragraph_format.line_spacing = None
        if para.text != "":
            para.add_run("\n\n")
        para.add_run(jst_outcome)
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        self.__Add_Checkbox(jst_outcome, percent)

    # adds both the Olivet course outcomes with there coresponding military course outcomes.
    # c_outcomes would be just a string for the Olivet college outcome and then jst_outcomes would
    # be the coresponding dictionary that holds percentages.
    def Like_Outcomes(self, c_outcome, jst_outcome):
        self.Olivet_Course_Outcomes(c_outcome)
        for outcome in jst_outcome:
            self.JST_Outcomes(outcome, jst_outcome[outcome]) 
        self.mc_row += 1

    # will be used if we decide to implement emailing the form.
    def Email_Doc(self):
        pass

    # used to save the document. Must call this to save document.
    def Save_Doc(self):
        self.doc.save('Test-Saved.docx')
